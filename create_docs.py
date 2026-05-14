"""Run this once to generate DOCUMENTATION.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json, os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

GOLD   = RGBColor(0xB8, 0x97, 0x2E)
BLACK  = RGBColor(0x11, 0x13, 0x18)
GREY   = RGBColor(0x55, 0x55, 0x66)
GREEN  = RGBColor(0x1A, 0x7A, 0x3C)
RED    = RGBColor(0x8B, 0x00, 0x00)

def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = GOLD
        run.font.size = Pt(18)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = BLACK
        run.font.size = Pt(13)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = GREY
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    if bold: run.bold = True
    if color: run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10.5)
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F4F4F4')
    p._element.get_or_add_pPr().append(shading)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1A1D24')
        cell._element.get_or_add_tcPr().append(shading)
    for r_idx, row_data in enumerate(rows):
        row = t.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if (r_idx % 2) == 0:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'F9F9F9')
                cell._element.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def divider():
    p = doc.add_paragraph('─' * 80)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC,0xCC,0xCC)
        run.font.size = Pt(8)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_heading('', 0)
title_p.clear()
run = title_p.add_run('Pearson Specter Litt')
run.font.color.rgb = GOLD
run.font.size = Pt(28)
run.bold = True
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Legal Document Processing Pipeline')
r.font.size = Pt(16)
r.font.color.rgb = BLACK

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub2.add_run('Complete Technical Documentation')
r2.font.size = Pt(12)
r2.font.color.rgb = GREY

doc.add_paragraph()
divider()

# ══════════════════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1('1. System Overview')

body(
    'Pearson Specter Litt is a four-stage AI pipeline that processes messy legal '
    'documents, retrieves grounded evidence, generates structured draft summaries, '
    'and continuously learns from operator edits to improve all future output.'
)

body('The four stages are:', bold=True)
bullet('Document Processing', bold_prefix='Stage 1 — ')
bullet('Grounded Retrieval', bold_prefix='Stage 2 — ')
bullet('Draft Generation', bold_prefix='Stage 3 — ')
bullet('Edit Capture and Pattern Learning', bold_prefix='Stage 4 — ')

body(
    'Every fact in the generated draft is pinned to a specific evidence chunk ID '
    'from the source document. The system explicitly lists what it could NOT find '
    'rather than hallucinating. Operator edits are analysed by Claude to extract '
    'transferable style rules that improve all future drafts automatically.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
h1('2. System Requirements')

h2('Minimum Requirements')
add_table(
    ['Component', 'Requirement'],
    [
        ['Python', '3.11 or higher'],
        ['RAM', '4 GB minimum (8 GB recommended)'],
        ['Disk', '2 GB free (models + ChromaDB)'],
        ['Network', 'Internet access for Anthropic API calls'],
        ['Anthropic API key', 'Required — get one at console.anthropic.com'],
    ],
    col_widths=[2.0, 4.0]
)

h2('Python Dependencies')
add_table(
    ['Package', 'Version', 'Purpose'],
    [
        ['anthropic', '>=0.30.0', 'Claude API for extraction, drafting, pattern learning'],
        ['pdfplumber', '>=0.10.0', 'Native PDF text extraction'],
        ['python-docx', '>=1.1.0', 'DOCX file extraction'],
        ['python-pptx', '>=0.6.23', 'PPTX / PPT file extraction'],
        ['Pillow', '>=10.0.0', 'Image file handling'],
        ['chromadb', '>=0.5.0', 'Persistent vector store'],
        ['sentence-transformers', '>=2.7.0', 'Text embeddings (all-MiniLM-L6-v2)'],
        ['fastapi', '>=0.111.0', 'REST API framework'],
        ['uvicorn', '>=0.30.0', 'ASGI server'],
        ['python-dotenv', '>=1.0.0', 'Environment variable loading'],
        ['pydantic', '>=2.0.0', 'Request/response validation'],
        ['aiofiles', '>=23.0.0', 'Static file serving'],
    ],
    col_widths=[2.0, 1.4, 2.8]
)

h2('Optional — OCR Support (for scanned PDFs and images)')
add_table(
    ['Package / Tool', 'Purpose'],
    [
        ['pytesseract', 'Python wrapper for Tesseract OCR'],
        ['pdf2image', 'Convert PDF pages to images for OCR'],
        ['Tesseract binary', 'OCR engine (system install required)'],
        ['poppler', 'PDF rendering library required by pdf2image'],
    ],
    col_widths=[2.5, 3.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. SETUP AND INSTALLATION
# ══════════════════════════════════════════════════════════════════════════════
h1('3. Setup and Installation')

h2('Step 1 — Create a virtual environment')
code_block('cd pearson-specter-litt\npython -m venv venv')

h2('Step 2 — Activate the virtual environment')
body('Windows:', bold=True)
code_block('venv\\Scripts\\activate')
body('macOS / Linux:', bold=True)
code_block('source venv/bin/activate')

h2('Step 3 — Install dependencies')
code_block('pip install -r requirements.txt')
body(
    'Note: On first run, sentence-transformers downloads the all-MiniLM-L6-v2 '
    'model (~90 MB). This happens once and is cached in your home directory.'
)

h2('Step 4 — Configure your API key')
code_block('# Windows\ncopy .env.example .env\n\n# macOS / Linux\ncp .env.example .env')
body('Open .env and set your Anthropic API key:')
code_block('ANTHROPIC_API_KEY=sk-ant-api03-...')

h2('Step 5 — Install OCR dependencies (optional)')
body('macOS:', bold=True)
code_block('brew install tesseract poppler\npip install pytesseract pdf2image')
body('Windows:', bold=True)
code_block(
    '1. Install Tesseract: https://github.com/UB-Mannheim/tesseract/wiki\n'
    '2. Add Tesseract to your PATH\n'
    '3. Download poppler: https://github.com/oschwartz10612/poppler-windows\n'
    '4. Add poppler bin/ to your PATH\n'
    '5. pip install pytesseract pdf2image'
)
body('Ubuntu / Debian:', bold=True)
code_block('sudo apt install tesseract-ocr poppler-utils\npip install pytesseract pdf2image')

# ══════════════════════════════════════════════════════════════════════════════
# 4. HOW TO RUN
# ══════════════════════════════════════════════════════════════════════════════
h1('4. How to Run')

h2('Option A — CLI Demo (recommended for reviewers)')
body(
    'The demo runs all four pipeline stages end-to-end on the three included '
    'sample documents and saves output files to sample_outputs/.'
)
code_block('python run_demo.py')

body('What happens when you run it:')
bullet('Stage 1: Ingests case_001_complaint.txt, case_002_contract.txt, case_003_notice.txt')
bullet('Stage 2: Indexes all chunks into ChromaDB with cosine similarity embeddings')
bullet('Stage 3: Retrieves top-5 evidence chunks and generates Case Fact Summary for case001')
bullet('Stage 4: Simulates 5 operator edits, extracts patterns, regenerates improved draft')
bullet('Saves 3 output files to sample_outputs/')

h2('Option B — API Server')
code_block('uvicorn api.app:app --reload')
body('Then open http://localhost:8000/docs for the interactive Swagger UI.')

h2('Option C — Web UI')
code_block('uvicorn api.app:app --reload')
body('Then open http://localhost:8000 for the full visual interface.')

h2('Option D — Docker')
code_block('docker compose up --build')
body('API and UI available at http://localhost:8000. ChromaDB data persists across restarts.')

h2('Running Tests')
code_block(
    '# Unit tests only — no API key needed (~10 seconds)\n'
    'pytest tests/ -v -m "not api"\n\n'
    '# All tests including Claude API integration\n'
    'pytest tests/ -v'
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. SUPPORTED INPUT FILE TYPES
# ══════════════════════════════════════════════════════════════════════════════
h1('5. Supported Input File Types')

add_table(
    ['Extension', 'Method', 'OCR Required?', 'Notes'],
    [
        ['.txt', 'Direct UTF-8 read', 'No', 'Most reliable — no conversion needed'],
        ['.pdf (text)', 'pdfplumber native', 'No', 'Auto-fallback to OCR if sparse'],
        ['.pdf (scanned)', 'pytesseract via pdf2image', 'Yes', 'Falls back gracefully if OCR unavailable'],
        ['.docx', 'python-docx', 'No', 'Extracts paragraphs + table cells'],
        ['.pptx / .ppt', 'python-pptx', 'No', 'Per-slide text labelled [Slide N]'],
        ['.png / .jpg / .jpeg', 'pytesseract', 'Yes', 'Direct image OCR at 300 DPI'],
        ['.tiff / .bmp', 'pytesseract', 'Yes', 'Same as PNG/JPG'],
    ],
    col_widths=[1.4, 1.8, 1.4, 2.2]
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. SAMPLE INPUTS
# ══════════════════════════════════════════════════════════════════════════════
h1('6. Sample Inputs')

add_table(
    ['File', 'Document Type', 'Notable Features'],
    [
        ['case_001_complaint.txt', 'Employment dispute complaint',
         'OCR artifact (Apri1), water damage notation, redacted address, missing exhibits A/B/C'],
        ['case_002_contract.txt', 'Software services agreement',
         'Payment terms, SLA clauses, indemnification provisions, defined terms'],
        ['case_003_notice.txt', 'Legal demand / notice letter',
         'Short form, partial jurisdiction info, demand deadline'],
    ],
    col_widths=[2.0, 1.8, 2.9]
)

body('All three documents are synthetic but realistic legal text designed to exercise:')
bullet('OCR artifact handling (typos, garbled characters from scanning)')
bullet('Partially missing or redacted information')
bullet('References to exhibits that are not present in the document')
bullet('Inconsistent formatting and date representations')

# ══════════════════════════════════════════════════════════════════════════════
# 7. SAMPLE OUTPUTS
# ══════════════════════════════════════════════════════════════════════════════
h1('7. Sample Outputs')

h2('Stage 1 Output — Structured Extraction')
body('After processing case_001_complaint.txt, the processor returns:')
code_block('''{
  "doc_id": "case001",
  "extraction_method": "text",
  "chunk_count": 11,
  "structured_fields": {
    "document_type": "complaint",
    "case_title": "Marcus A. Hendricks v. Nexigen Pharmaceuticals, Inc., et al.",
    "case_number": "24-CV-08821",
    "parties": [
      { "name": "Marcus A. Hendricks", "role": "plaintiff" },
      { "name": "Nexigen Pharmaceuticals, Inc.", "role": "defendant" },
      { "name": "Raymond L. Porter", "role": "defendant" }
    ],
    "legal_issues": ["Breach of Contract", "Fraud and Deceit",
                     "Negligent Misrepresentation",
                     "California Business and Professions Code Section 17200"],
    "jurisdiction": "Superior Court of California, County of Los Angeles",
    "document_quality": "OCR artifact Apri1 on page 3; water damage noted upper-left page 2"
  },
  "processing_notes": [
    "Produced 11 chunks (method=text)"
  ]
}''')

h2('Stage 2+3 Output — Case Fact Summary (v1, no patterns)')
body('File: sample_outputs/case001_draft_v1.json')
code_block('''{
  "case_title": "Marcus A. Hendricks v. Nexigen Pharmaceuticals, Inc., et al.",
  "case_number": "24-CV-08821",
  "document_type": "complaint",
  "jurisdiction": "Superior Court of the State of California, County of Los Angeles",
  "parties": [
    {
      "name": "Marcus A. Hendricks",
      "role": "Plaintiff",
      "evidence_id": "case001__chunk_0"
    },
    {
      "name": "Nexigen Pharmaceuticals, Inc.",
      "role": "Defendant (Delaware Corporation)",
      "evidence_id": "case001__chunk_0"
    }
  ],
  "key_facts": [
    {
      "fact": "Hendricks was employed as Senior Director of Business Development at
               Nexigen from approximately April 2019 through October 14, 2023.",
      "evidence_id": "case001__chunk_2",
      "evidence_quote": "Hendricks was employed as Senior Director ... from Apri1 2019
                         through October 14, 2023"
    },
    {
      "fact": "Total claimed damages are not less than $2,704,125 comprising $210,000
               severance and $2,494,125 RSU value.",
      "evidence_id": "case001__chunk_9",
      "evidence_quote": "damages not less than $2,704,125 (comprising $210,000 severance
                         + $2,494,125 RSU value)"
    }
  ],
  "missing_information": [
    "Exhibits A, B, and C referenced but not attached — Employment Agreement absent.",
    "RSU valuation of $22.17/share from partially illegible page — requires verification.",
    "Raymond L. Porter residential address redacted — service of process unverifiable.",
    "Case number prefix 24 inconsistent with November 2023 filing date — unresolved."
  ],
  "analyst_notes": "First-pass review based on five retrieved evidence passages."
}''')

h2('Stage 3 Output — Evidence Used')
body('Alongside the draft, the system returns the retrieved evidence passages:')
code_block('''"evidence_used": [
  {
    "chunk_id": "case001__chunk_0",
    "score": 0.4105,
    "text_preview": "SUPERIOR COURT OF THE STATE OF CALIFORNIA ... BREACH OF CONTRACT ..."
  },
  {
    "chunk_id": "case001__chunk_9",
    "score": 0.3164,
    "text_preview": "damages not less than $2,704,125 (comprising $210,000 severance + ...)"
  }
]''')

h2('Stage 4 Output — Edit Record and Learned Patterns')
body('File: sample_outputs/case001_edit_record.json')
code_block('''{
  "edit_id": "a3f9c12e8b04",
  "doc_id": "case001",
  "timestamp": "2026-05-14T10:31:47Z",
  "operator_edits_applied": [
    "Added firm-standard case_reference header field",
    "Normalised dates to ISO format (YYYY-MM-DD)",
    "Wrapped evidence_quote values in typographic quotation marks",
    "Trimmed speculative analyst_notes to first sentence only",
    "Promoted jurisdiction to top-level filing_jurisdiction field"
  ],
  "patterns_extracted": [
    {
      "description": "ISO-format all dates",
      "rule": "Normalise all date values in relevant_dates to ISO 8601 (YYYY-MM-DD)",
      "example_before": "\\"date\\": \\"11/08/2023\\"",
      "example_after": "\\"date\\": \\"2023-11-08\\"",
      "confidence": 0.98
    },
    {
      "description": "Trim speculative analyst notes",
      "rule": "Limit analyst_notes to a single declarative sentence",
      "example_before": "This summary ... should be verified before litigation.",
      "example_after": "This summary is a first-pass analysis.",
      "confidence": 0.88
    }
  ],
  "total_active_patterns": 5
}''')

h2('Before and After — Draft Improvement')
add_table(
    ['Field', 'v1 (no patterns)', 'v2 (patterns applied)'],
    [
        ['analyst_notes',
         'This summary is a first-pass analysis based solely on five retrieved evidence passages ... should be verified against the original complaint before any reliance for litigation purposes.',
         'This summary is a first-pass analysis based solely on five retrieved evidence passages.'],
        ['evidence_quote style',
         'Hendricks was employed as Senior Director...',
         '"Hendricks was employed as Senior Director..."'],
        ['date format',
         '2019-04 (approximate)',
         '2019-04 (approx.) — ISO format enforced for all slash-dates'],
        ['Top-level fields',
         'case_title, case_number, document_type ...',
         'Adds case_reference and filing_jurisdiction at top level'],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. API REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
h1('8. API Reference')
body('Base URL: http://localhost:8000   |   Interactive docs: http://localhost:8000/docs')

h2('Upload and Process a Document')
code_block('POST /documents/upload\n\ncurl -X POST http://localhost:8000/documents/upload \\\n  -F "file=@case_001_complaint.txt" \\\n  -F "doc_id=case001"')
body('Response: doc_id, extraction_method, chunk_count, structured_fields, processing_notes')

h2('Generate a Grounded Draft')
code_block('GET /documents/{doc_id}/draft\n\ncurl http://localhost:8000/documents/case001/draft')
body('Response: draft JSON with key_facts, parties, dates, legal_issues, missing_information, evidence_used')

h2('Submit an Operator Edit')
code_block('POST /edits\n\ncurl -X POST http://localhost:8000/edits \\\n  -H "Content-Type: application/json" \\\n  -d \'{"doc_id":"case001","original_draft":{...},"edited_draft":{...}}\'')
body('Response: edit_id, patterns_extracted, total_active_patterns')

h2('Other Endpoints')
add_table(
    ['Method', 'Path', 'Description'],
    [
        ['GET', '/patterns', 'List all active learned patterns'],
        ['POST', '/patterns/deactivate', 'Suppress a pattern by pattern_id'],
        ['GET', '/documents/{doc_id}/edits', 'Edit history for a document'],
        ['GET', '/health', 'Returns {"status": "ok"}'],
        ['GET', '/docs', 'Interactive Swagger UI'],
    ],
    col_widths=[0.8, 2.2, 3.2]
)

# ══════════════════════════════════════════════════════════════════════════════
# 9. PROJECT FILE STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
h1('9. Project File Structure')
code_block(
    'pearson-specter-litt/\n'
    '├── src/\n'
    '│   ├── config.py          API key, model, chunk sizes, paths\n'
    '│   ├── models.py          Data classes: ProcessedDocument, DraftOutput ...\n'
    '│   ├── processor.py       Stage 1: extraction + OCR + structured fields\n'
    '│   ├── retriever.py       Stage 2: ChromaDB vector store + query\n'
    '│   ├── drafter.py         Stage 3: grounded Case Fact Summary via Claude\n'
    '│   ├── editor.py          Stage 4: edit capture + SQLite pattern store\n'
    '│   └── pipeline.py        Orchestrator tying all stages together\n'
    '├── api/\n'
    '│   └── app.py             FastAPI REST interface (7 endpoints)\n'
    '├── ui/\n'
    '│   └── index.html         Web UI (no build step required)\n'
    '├── sample_documents/\n'
    '│   ├── case_001_complaint.txt\n'
    '│   ├── case_002_contract.txt\n'
    '│   └── case_003_notice.txt\n'
    '├── sample_outputs/\n'
    '│   ├── case001_draft_v1.json\n'
    '│   ├── case001_draft_v2.json\n'
    '│   └── case001_edit_record.json\n'
    '├── tests/\n'
    '│   └── test_pipeline.py   Unit + integration tests\n'
    '├── run_demo.py            End-to-end CLI demo\n'
    '├── requirements.txt\n'
    '├── Dockerfile\n'
    '└── docker-compose.yml'
)

# ══════════════════════════════════════════════════════════════════════════════
# 10. ASSUMPTIONS AND TRADEOFFS
# ══════════════════════════════════════════════════════════════════════════════
h1('10. Assumptions and Tradeoffs')

add_table(
    ['Decision', 'Chosen Approach', 'Alternative', 'Reason'],
    [
        ['Vector DB', 'ChromaDB local', 'Pinecone, Weaviate', 'Zero infra; swap is one constructor change'],
        ['Embeddings', 'sentence-transformers CPU', 'OpenAI text-embedding-3-small', 'No extra API key; offline capable'],
        ['Pattern storage', 'SQLite', 'PostgreSQL', 'Single-file deployment; good for prototype'],
        ['OCR engine', 'pytesseract', 'AWS Textract, Azure OCR', 'Open-source; no cloud dependency'],
        ['Draft format', 'Strict JSON schema', 'Free-form prose', 'Machine-verifiable grounding; parseable'],
        ['LLM', 'Claude (all stages)', 'Separate specialised models', 'Single API key; consistent quality'],
    ],
    col_widths=[1.3, 1.6, 1.6, 2.0]
)

h2('Known Limitations')
bullet('Session cache: pipeline._docs is in-memory. After API restart, documents in ChromaDB cannot be drafted without re-ingesting.')
bullet('No cross-document retrieval: each draft is anchored to a single doc_id.')
bullet('Pattern deduplication is string-based: semantically identical patterns with different descriptions are stored separately.')
bullet('No streaming: API returns complete JSON responses.')

# ══════════════════════════════════════════════════════════════════════════════
# 11. EVALUATION RESULTS
# ══════════════════════════════════════════════════════════════════════════════
h1('11. Evaluation Results')

h2('Stage 1 — Document Processing Coverage')
add_table(
    ['Document', 'Method', 'Chunks', 'Fields Populated'],
    [
        ['case_001_complaint.txt', 'text', '11', '6 / 6 — 100%'],
        ['case_002_contract.txt', 'text', '8', '6 / 6 — 100%'],
        ['case_003_notice.txt', 'text', '5', '5 / 6 — 83% (jurisdiction partial)'],
    ],
    col_widths=[2.5, 1.2, 0.8, 1.8]
)

h2('Stage 2 — Retrieval Spot-Check (case001)')
add_table(
    ['Rank', 'Chunk ID', 'Score', 'Content Preview'],
    [
        ['1', 'case001__chunk_0', '0.41', 'Court header, parties, claims list'],
        ['2', 'case001__chunk_1', '0.40', 'Jurisdiction and venue allegations'],
        ['3', 'case001__chunk_10', '0.37', 'Counsel, filing stamp, missing exhibits'],
        ['4', 'case001__chunk_9', '0.32', 'Damages figure ($2,704,125), prayer for relief'],
        ['5', 'case001__chunk_2', '0.24', 'Party descriptions, employment dates'],
    ],
    col_widths=[0.6, 1.8, 0.8, 3.0]
)

h2('Stage 4 — Pattern Adoption Rate')
add_table(
    ['Pattern', 'Adopted in v2?'],
    [
        ['Add case_reference header field', 'Yes — field present in v2'],
        ['Normalise dates to ISO 8601', 'Yes — all dates in YYYY-MM-DD'],
        ['Wrap evidence_quote in double-quotes', 'Yes — quotes wrapped'],
        ['Trim analyst_notes to one sentence', 'Yes — single sentence only'],
        ['Promote jurisdiction to filing_jurisdiction', 'Yes — top-level field added'],
    ],
    col_widths=[3.5, 2.5]
)
body('Result: 5 / 5 patterns adopted — 100% pattern adoption rate on simulated edits.', bold=True)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = 'DOCUMENTATION.docx'
doc.save(out)
print(f'Saved: {out}')
